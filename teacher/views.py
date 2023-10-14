from django.shortcuts import render,redirect,reverse
from . import forms,models
from django.db.models import Sum
from django.contrib.auth.models import Group
from django.http import HttpResponseRedirect
from django.contrib.auth.decorators import login_required,user_passes_test
from django.conf import settings
from datetime import date, timedelta
from exam import models as QMODEL
from student import models as SMODEL
from exam import forms as QFORM

from http.client import HTTPResponse
import os
from django.core.files.storage import FileSystemStorage
import pandas as pd
import json
import sys
sys.path.insert(1,'D:\\Dev\\onlinexamination')

#for showing signup/login button for teacher
def teacherclick_view(request):
    if request.user.is_authenticated:
        return HttpResponseRedirect('afterlogin')
    return render(request,'teacher/teacherclick.html')

def teacher_signup_view(request):
    userForm=forms.TeacherUserForm()
    teacherForm=forms.TeacherForm()
    mydict={'userForm':userForm,'teacherForm':teacherForm}
    if request.method=='POST':
        userForm=forms.TeacherUserForm(request.POST)
        teacherForm=forms.TeacherForm(request.POST,request.FILES)
        if userForm.is_valid() and teacherForm.is_valid():
            user=userForm.save()
            user.set_password(user.password)
            user.save()
            teacher=teacherForm.save(commit=False)
            teacher.user=user
            teacher.save()
            my_teacher_group = Group.objects.get_or_create(name='TEACHER')
            my_teacher_group[0].user_set.add(user)
        return HttpResponseRedirect('teacherlogin')
    return render(request,'teacher/teachersignup.html',context=mydict)


def is_teacher(user):
    return user.groups.filter(name='TEACHER').exists()


# if QMODEL.Course.objects.filter(teacher =models.Teacher.objects.get(user_id=request.user.id)).exists() else 0
@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def teacher_dashboard_view(request):
    teacher = models.Teacher.objects.get(user_id=request.user.id)
    course=QMODEL.Course.objects.all().filter(teacher=teacher)
    num_of_Q=0
    for x in course:
        num_of_Q+=QMODEL.Question.objects.all().filter(course_id=x).count()

    dict={
    'total_course':course.count(),
    'total_question':num_of_Q,
    'total_student':SMODEL.Student.objects.all().count()
    }
    return render(request,'teacher/teacher_dashboard.html',context=dict)


@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def teacher_exam_view(request):
    return render(request,'teacher/teacher_exam.html')



@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def teacher_add_exam_view(request):
    courseForm=QFORM.CourseForm()
    if request.method=='POST':
        courseForm=QFORM.CourseForm(request.POST)
        if courseForm.is_valid():
            courseForm = courseForm.save(commit=False)
            courseForm.teacher = models.Teacher.objects.get(user_id = request.user.id)
            courseForm.save()
        else:
            print("form is invalid")
        return HttpResponseRedirect('/teacher/teacher-view-exam')
    return render(request,'teacher/teacher_add_exam.html',{'courseForm':courseForm})


@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def teacher_view_exam_view(request):
    # courses = QMODEL.Course.objects.all()
    courses = QMODEL.Course.objects.filter(teacher =models.Teacher.objects.get(user_id=request.user.id))
    return render(request,'teacher/teacher_view_exam.html',{'courses':courses})


@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def delete_exam_view(request,pk):
    course=QMODEL.Course.objects.get(id=pk)
    course.delete()
    return HttpResponseRedirect('/teacher/teacher-view-exam')


@login_required(login_url='adminlogin')
def teacher_question_view(request):
    
    return render(request,'teacher/teacher_question.html')


@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def teacher_add_question_view(request):

    teacher = models.Teacher.objects.get(user_id=request.user.id)
    custom_queryset = QMODEL.Course.objects.filter(teacher=teacher)

    questionForm=QFORM.QuestionForm(custom_queryset=custom_queryset)
    if request.method=='POST':
        questionForm=QFORM.QuestionForm(request.POST)
        if questionForm.is_valid():
            question=questionForm.save(commit=False)
            course=QMODEL.Course.objects.get(id=request.POST.get('courseID'))
            question.course=course
            question.save()       
        else:
            print("form is invalid")
        return HttpResponseRedirect('/teacher/teacher-view-question')
    return render(request,'teacher/teacher_add_question.html',{'questionForm':questionForm})


@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def teacher_view_question_view(request):
    teacher = models.Teacher.objects.get(user_id=request.user.id)
    courses= QMODEL.Course.objects.filter(teacher=teacher )
    return render(request,'teacher/teacher_view_question.html',{'courses':courses})


@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def see_question_view(request,pk):
    request.session['course-id']=pk
    questions=QMODEL.Question.objects.all().filter(course_id=pk)
    number_of_books = QMODEL.PDFDocument.objects.all().count()

    data_to_export = {'pdf_path': ""}
    serialized_data = json.dumps(data_to_export)
    with open('data.json', 'w') as file:
        file.write(serialized_data)

    if request.method == 'POST':  
        
        QMODEL.PDFDocument.objects.create(id=number_of_books+1, title= request.POST["title"], pdf_file=request.FILES['pdf_file'],
                                          teacher= models.Teacher.objects.get(user_id = request.user.id),
                                        course=QMODEL.Course.objects.get(id = pk)).save()
        
        obj = QMODEL.PDFDocument.objects.get(id=number_of_books+1)
        pdf_path = obj.pdf_file.path

        data_to_export = {'pdf_path': pdf_path}
        serialized_data = json.dumps(data_to_export)

        with open('data.json', 'w') as file:
            file.write(serialized_data)
        
        check1=request.POST.get('easy',0)
        check2=request.POST.get('mid',0)
        check3=request.POST.get('hard',0)
        check4=request.POST.get('veryhard',0)

        MCQ=request.POST.get('multichoice',0)
        TF=request.POST.get('truefalse',0)

        import untitled139
        from untitled139 import d, multichoicequestiondict, multichoicequestiondict2, multichoicequestiondict3, multichoicequestiondict1
        
        if check1 == "1" and MCQ:
            i = 0
            j = 0
            for choices in multichoicequestiondict3['multichoice']:
                if multichoicequestiondict3['answer'][i] == choices[j]:
                    ans = 'Option1'
                elif multichoicequestiondict3['answer'][i] == choices[j+1]:
                    ans = 'Option2'
                elif multichoicequestiondict3['answer'][i] == choices[j+2]:
                    ans = 'Option3'
                else :
                    ans = 'Option4'
                QMODEL.Question.objects.create(rank=1,question=multichoicequestiondict3['question'][i], option1=choices[j], option2=choices[j+1], option3=choices[j+2], option4=choices[j+3], answer=ans, course_id = pk).save()
                i+=1
        
        if check2 =="2" and MCQ:
            i = 0
            j = 0
            for choices in multichoicequestiondict2['multichoice']:
                if multichoicequestiondict2['answer'][i] == choices[j]:
                    ans = 'Option1'
                elif multichoicequestiondict2['answer'][i] == choices[j+1]:
                    ans = 'Option2'
                elif multichoicequestiondict2['answer'][i] == choices[j+2]:
                    ans = 'Option3'
                else :
                    ans = 'Option4'
                QMODEL.Question.objects.create(marks=2,rank=2,question=multichoicequestiondict2['question'][i], option1=choices[j], option2=choices[j+1], option3=choices[j+2], option4=choices[j+3], answer=ans, course_id = pk).save()
                i+=1

        if check3 =="3" and MCQ:
            i = 0
            j = 0
            for choices in multichoicequestiondict1['multichoice']:
                if multichoicequestiondict1['answer'][i] == choices[j]:
                    ans = 'Option1'
                elif multichoicequestiondict1['answer'][i] == choices[j+1]:
                    ans = 'Option2'
                elif multichoicequestiondict1['answer'][i] == choices[j+2]:
                    ans = 'Option3'
                else :
                    ans = 'Option4'
                QMODEL.Question.objects.create(marks=3,rank=3,question=multichoicequestiondict1['question'][i], option1=choices[j], option2=choices[j+1], option3=choices[j+2], option4=choices[j+3], answer=ans, course_id = pk).save()
                i+=1

        if check4 =="4" and MCQ:
            i = 0
            j = 0
            for choices in multichoicequestiondict['multichoice']:
                if multichoicequestiondict['answer'][i] == choices[j]:
                    ans = 'Option1'
                elif multichoicequestiondict['answer'][i] == choices[j+1]:
                    ans = 'Option2'
                elif multichoicequestiondict['answer'][i] == choices[j+2]:
                    ans = 'Option3'
                else :
                    ans = 'Option4'
                QMODEL.Question.objects.create(marks=3,rank=4,question=multichoicequestiondict['question'][i], option1=choices[j], option2=choices[j+1], option3=choices[j+2], option4=choices[j+3], answer=ans, course_id = pk).save()
                i+=1
        #######################################################
        if TF :
            i = 0
            j = 0
            for choices in d['multichoice']:
                if d['answer'][i] == choices[j]:
                    ans = 'Option1'
                elif d['answer'][i] == choices[j+1]:
                    ans = 'Option2'
                QMODEL.Question.objects.create(question=d['question'][i], option1=choices[j], option2=choices[j+1], answer=ans, course_id = pk).save()
                i+=1
                
    return render(request,'teacher/see_question.html',{'questions':questions })
    

@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def remove_question_view(request,pk):
    question=QMODEL.Question.objects.get(id=pk)
    course_id=request.session['course-id']
    question.delete()
    return HttpResponseRedirect('/teacher/see-question/'+str(course_id))


@login_required(login_url='teacherlogin')
@user_passes_test(is_teacher)
def update_question_view(request,pk):

    question=QMODEL.Question.objects.get(id=pk)
    course_id=request.session['course-id']
    if request.method=='POST':
        question.id = pk
        question.question = request.POST["Question"]
        question.marks = request.POST["Marks"]
        question.option1 = request.POST["Option1"]
        question.option2 = request.POST["Option2"]
        question.option3 = request.POST["Option3"]
        question.option4 = request.POST["Option4"]
        question.rank = request.POST["rank"]
        question.answer = request.POST["options"]
        question.save()

        return HttpResponseRedirect('/teacher/see-question/'+str(course_id))

    return render(request,'teacher/update_question.html',{'question':question})


import docx
def questions_Toword(request):
    temp = {}
    questions_list=[]
    answers_list =[]
    list_of_choices=[]

    pk=request.session['course-id']
    questions=QMODEL.Question.objects.all().filter(course_id=pk)

    for obj in questions:
        temp_list=[]
        questions_list.append(obj.question)
        answers_list.append(obj.answer)
        temp_list.append(obj.option1)
        temp_list.append(obj.option2)
        temp_list.append(obj.option3)
        temp_list.append(obj.option4)
        list_of_choices.append(temp_list)

    # for x in answers_list:
    #     if x == 'Option1'
           
        

    temp['question']=questions_list
    temp['answer']=answers_list
    temp['multichoice']=list_of_choices

    doc = docx.Document()
    for i, question in enumerate(temp["question"]):
        doc.add_paragraph(f"Q{i+1}: {question}")
        for choice in temp['multichoice'][i]:
            if choice == None:
                continue

            doc.add_paragraph(str(choice), style='List Bullet')
        doc.add_paragraph(f"Answer: {temp['answer'][i]}")
    doc.save("cyber_security_questions.docx")

    return HttpResponseRedirect('/teacher/see-question/'+str(pk))

########################################################################################## temporRY


@login_required(login_url='teacherlogin')
def teacher_view_student_marks_view(request):
    students= SMODEL.Student.objects.all()
    return render(request,'teacher/teacher_view_student_marks.html',{'students':students})

@login_required(login_url='teacherlogin')
def teacher_view_marks_view(request,pk):
    teacher = models.Teacher.objects.get(user_id=request.user.id)
    courses= QMODEL.Course.objects.filter(teacher=teacher )
    response =  render(request,'teacher/teacher_view_marks.html',{'courses':courses})
    response.set_cookie('student_id',str(pk))
    return response

@login_required(login_url='teacherlogin')
def teacher_check_marks_view(request,pk):
    course = QMODEL.Course.objects.get(id=pk)
    student_id = request.COOKIES.get('student_id')
    student= SMODEL.Student.objects.get(id=student_id)

    results= QMODEL.Result.objects.all().filter(exam=course).filter(student=student)
    return render(request,'teacher/teacher_check_marks.html',{'results':results})
    

